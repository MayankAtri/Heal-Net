#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('HealNet - Medical Assistant Application', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Comprehensive Presentation Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 1. PROJECT OVERVIEW
doc.add_heading('1. PROJECT OVERVIEW', 1)
doc.add_paragraph('HealNet is a comprehensive AI-powered medical assistant web application that helps users understand their prescriptions, medical reports, and over-the-counter medications.')

doc.add_heading('Key Value Proposition:', 2)
points = [
    'Simplifies complex medical jargon into easy-to-understand language',
    'Uses AI (Google Gemini) and OCR (Tesseract.js) for intelligent analysis',
    'Provides medicine disposal information for safe healthcare practices',
    'Maintains complete history of all medical analyses',
    'Supports both authenticated and guest users'
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

# 2. CORE FEATURES
doc.add_heading('2. CORE FEATURES', 1)

doc.add_heading('2.1 Prescription Analysis 💊', 2)
features = [
    'Upload prescription images (JPG, PNG, JPEG)',
    'OCR text extraction using Tesseract.js',
    'AI-powered simplification using Google Gemini Vision API',
    'Detailed analysis: medicine names, dosages, frequency, duration, warnings',
    'Storage and retrieval of prescription history',
    'Disposal information for safe medicine disposal'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.2 Medical Reports Analysis 📊', 2)
features = [
    'Analyze blood tests, X-rays, MRI scans, pathology reports',
    'Customizable analysis depth (Simple, Detailed, Educational)',
    'AI-powered interpretation of medical results',
    'Easy-to-understand explanations',
    'Possible conditions identification with likelihood assessment'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3 OTC Consultation 🏥', 2)
features = [
    'Over-the-counter medicine suggestions for 14+ common symptoms',
    'Age-specific recommendations (infant, child, teen, adult, senior)',
    'Safety information and warnings',
    'Dosage recommendations with practical instructions',
    'Side effects information',
    'Home remedies and when to see a doctor'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.4 Medicine Disposal Info ♻️', 2)
features = [
    'Interactive map showing nearby disposal locations',
    'Uses OpenStreetMap, Leaflet.js, and Overpass API',
    'Geolocation to find user\'s current location',
    'Distance calculation to nearest disposal points',
    'Comprehensive disposal guidelines for different medicine types',
    'Environmental impact information'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.5 User Authentication & History 👤', 2)
features = [
    'Google OAuth 2.0 authentication',
    'Email/password authentication with JWT',
    'User profile management',
    'Complete history of all analyses',
    'Detail view pages for each analysis type'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# 3. TECHNICAL ARCHITECTURE
doc.add_heading('3. TECHNICAL ARCHITECTURE', 1)

doc.add_heading('3.1 Frontend Stack', 2)
tech = [
    'React 19.2.0 - Modern UI framework',
    'Vite 7.2.4 - Fast build tool',
    'Tailwind CSS 3.4.19 - Utility-first styling',
    'Framer Motion 12.29.2 - Smooth animations',
    'React Router DOM 7.11.0 - Client-side routing',
    'Axios 1.13.2 - HTTP client',
    'Leaflet 1.9.4 & React-Leaflet 5.0.0 - Interactive maps',
    'React Dropzone 14.3.8 - File uploads',
    'Glassmorphism UI design - Modern aesthetic'
]
for t in tech:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('3.2 Backend Stack', 2)
tech = [
    'Node.js with Express.js 4.18.2',
    'MongoDB with Mongoose 8.0.3',
    'Google Gemini API 0.21.0 - AI analysis',
    'Tesseract.js 5.0.4 - OCR processing',
    'Passport.js - Authentication (Google OAuth & JWT)',
    'bcryptjs 3.0.3 - Password hashing',
    'Multer 1.4.5 - File upload handling',
    'Helmet 7.1.0 - Security headers',
    'Express Rate Limit 7.1.5 - API protection'
]
for t in tech:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('3.3 External APIs & Services', 2)
apis = [
    'Google Gemini AI - Medical text analysis',
    'Tesseract.js - OCR extraction',
    'OpenStreetMap - Free map tiles',
    'Overpass API - Query nearby pharmacies/hospitals',
    'Nominatim - Geocoding service',
    'Browser Geolocation API - User location'
]
for api in apis:
    doc.add_paragraph(api, style='List Bullet')

doc.add_page_break()

# 4. SYSTEM DESIGN
doc.add_heading('4. SYSTEM DESIGN', 1)

doc.add_heading('4.1 Architecture Pattern', 2)
patterns = [
    'Frontend: Component-based architecture with React',
    'Backend: MVC (Model-View-Controller) pattern',
    'API: RESTful API design',
    'Authentication: JWT with httpOnly cookies + OAuth 2.0',
    'State Management: React Context API (DarkModeContext, AuthContext)'
]
for p in patterns:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('4.2 Data Models', 2)
doc.add_paragraph('User Model:')
fields = ['name, email, password (hashed)', 'googleId, profilePicture', 'authProvider (local/google/both)', 'refreshTokens array', 'timestamps']
for f in fields:
    doc.add_paragraph(f, style='List Bullet 2')

doc.add_paragraph('Prescription Model:')
fields = ['originalText (OCR result)', 'simplifiedAnalysis (medicines, warnings, instructions)', 'imageUrl, originalFilename', 'userId reference', 'processingStatus']
for f in fields:
    doc.add_paragraph(f, style='List Bullet 2')

doc.add_paragraph('Medical Report Model:')
fields = ['reportType (blood_test/radiology/pathology)', 'analysisDepth (simple/detailed/educational)', 'analysis (polymorphic based on type)', 'bloodTestResults, radiologyFindings, pathologyFindings', 'possibleConditions with likelihood']
for f in fields:
    doc.add_paragraph(f, style='List Bullet 2')

doc.add_paragraph('OTC Consultation Model:')
fields = ['symptomType (14+ predefined types)', 'ageGroup (infant/child/teen/adult/senior)', 'suggestions (medicines, homeRemedies, whenToSeeDoctor)', 'medicalDisclaimer']
for f in fields:
    doc.add_paragraph(f, style='List Bullet 2')

doc.add_page_break()

# 5. KEY WORKFLOWS
doc.add_heading('5. KEY WORKFLOWS', 1)

doc.add_heading('5.1 Prescription Analysis Workflow', 2)
steps = [
    '1. User uploads prescription image via drag-and-drop',
    '2. Frontend validates file (type, size < 5MB)',
    '3. Backend receives image via Multer',
    '4. Gemini Vision API analyzes image directly (no OCR needed)',
    '5. AI extracts medicines, dosages, warnings, instructions',
    '6. Structured data saved to MongoDB',
    '7. Response sent to frontend with analysis',
    '8. User views simplified prescription with all details'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('5.2 Authentication Flow', 2)
steps = [
    '1. User chooses Google OAuth or Email/Password',
    '2. Google OAuth: Redirect to Google → Callback → JWT token',
    '3. Email/Password: Validate credentials → Hash password → JWT token',
    '4. JWT stored in httpOnly cookie + localStorage',
    '5. Protected routes check JWT validity',
    '6. Refresh token mechanism for session management'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# 6. SECURITY FEATURES
doc.add_heading('6. SECURITY FEATURES', 1)
security = [
    'Rate limiting (100 requests per 15 minutes per IP)',
    'CORS configuration for specific origins',
    'Helmet security headers',
    'JWT-based authentication with httpOnly cookies',
    'Password hashing with bcryptjs (10 salt rounds)',
    'File upload validation (type, size, content)',
    'Input validation with express-validator',
    'MongoDB injection prevention via Mongoose',
    'Trust proxy configuration for Render deployment'
]
for s in security:
    doc.add_paragraph(s, style='List Bullet')

# 7. DEPLOYMENT
doc.add_heading('7. DEPLOYMENT', 1)
deployment = [
    'Frontend: Vercel (with preview deployments)',
    'Backend: Render',
    'Database: MongoDB Atlas',
    'Environment: Production-ready with proper CORS',
    'Cross-domain authentication handling',
    'Environment variables for sensitive data'
]
for d in deployment:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# 8. PRESENTATION TALKING POINTS
doc.add_heading('8. PRESENTATION TALKING POINTS', 1)

doc.add_heading('Opening (1-2 minutes)', 2)
doc.add_paragraph('Problem Statement:', style='List Bullet')
doc.add_paragraph('Medical prescriptions and reports are often filled with complex jargon that patients struggle to understand. This leads to medication errors, improper disposal, and unnecessary anxiety.', style='List Bullet 2')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('HealNet uses AI to bridge the gap between medical professionals and patients, making healthcare information accessible to everyone.', style='List Bullet 2')

doc.add_heading('Demo Flow (5-7 minutes)', 2)
demo = [
    '1. Show homepage with glassmorphism design',
    '2. Upload a sample prescription → Show OCR + AI analysis',
    '3. Navigate to Medical Reports → Upload blood test → Show detailed analysis',
    '4. Go to OTC Consultation → Select "headache" → Show recommendations',
    '5. Open Medicine Disposal → Show map with nearby locations',
    '6. Show History page with all past analyses',
    '7. Demonstrate Google OAuth login'
]
for d in demo:
    doc.add_paragraph(d, style='List Number')

doc.add_heading('Technical Highlights (2-3 minutes)', 2)
highlights = [
    'AI-Powered: Google Gemini Vision API for direct image analysis',
    'Modern Stack: React 19, Vite, Tailwind CSS, Node.js, MongoDB',
    'Security: JWT authentication, rate limiting, input validation',
    'User Experience: Glassmorphism UI, smooth animations, responsive design',
    'Scalability: RESTful API, MongoDB indexing, optimized queries'
]
for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

doc.add_heading('Unique Features (1-2 minutes)', 2)
unique = [
    'Direct image analysis with Gemini Vision (no separate OCR step)',
    'Customizable analysis depth for medical reports',
    'Age-specific OTC recommendations',
    'Interactive medicine disposal map with real-time locations',
    'Complete history tracking for all analysis types',
    'Dark mode support throughout the application'
]
for u in unique:
    doc.add_paragraph(u, style='List Bullet')

doc.add_heading('Impact & Use Cases (1-2 minutes)', 2)
impact = [
    'Patients: Better understand prescriptions, reduce medication errors',
    'Elderly: Simplified medical information in plain language',
    'Parents: Safe OTC recommendations for children',
    'Environment: Proper medicine disposal reduces pollution',
    'Healthcare: Reduces burden on pharmacists for basic queries'
]
for i in impact:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('Future Enhancements (1 minute)', 2)
future = [
    'Multi-language support (Hindi, Spanish, etc.)',
    'Voice input for symptoms',
    'Drug interaction checker',
    'Medication reminders and tracking',
    'Integration with pharmacy APIs for price comparison',
    'Telemedicine consultation booking'
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# 9. Q&A PREPARATION
doc.add_heading('9. Q&A PREPARATION', 1)

doc.add_heading('Expected Questions & Answers:', 2)

doc.add_paragraph('Q: How accurate is the AI analysis?')
doc.add_paragraph('A: We use Google Gemini, a state-of-the-art AI model. However, we include medical disclaimers and recommend consulting healthcare professionals for critical decisions.', style='List Bullet 2')

doc.add_paragraph('Q: Is user data secure?')
doc.add_paragraph('A: Yes. We use JWT authentication, password hashing, CORS protection, rate limiting, and store data in MongoDB Atlas with encryption.', style='List Bullet 2')

doc.add_paragraph('Q: Can it handle handwritten prescriptions?')
doc.add_paragraph('A: Yes! Gemini Vision API is trained to read both printed and handwritten text, including doctor\'s handwriting.', style='List Bullet 2')

doc.add_paragraph('Q: What happens if the AI makes a mistake?')
doc.add_paragraph('A: We display clear medical disclaimers. The app is designed as an assistive tool, not a replacement for professional medical advice.', style='List Bullet 2')

doc.add_paragraph('Q: How do you handle different types of medical reports?')
doc.add_paragraph('A: We have polymorphic data models that adapt based on report type (blood test, radiology, pathology) with specialized analysis for each.', style='List Bullet 2')

doc.add_paragraph('Q: Is the application scalable?')
doc.add_paragraph('A: Yes. We use MongoDB for horizontal scaling, implement caching strategies, rate limiting, and the architecture supports microservices migration.', style='List Bullet 2')

doc.add_page_break()

# 10. TECHNICAL DEEP DIVE
doc.add_heading('10. TECHNICAL DEEP DIVE', 1)

doc.add_heading('API Endpoints:', 2)
endpoints = [
    'POST /api/auth/register - User registration',
    'POST /api/auth/login - User login',
    'GET /api/auth/google - Google OAuth',
    'POST /api/prescription/upload - Analyze prescription',
    'GET /api/prescription/:id - Get prescription details',
    'POST /api/reports/upload - Analyze medical report',
    'POST /api/otc/consult - Get OTC suggestions',
    'GET /api/history - Get user history'
]
for e in endpoints:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('Performance Optimizations:', 2)
optimizations = [
    'Reduced search radius for map queries (3km)',
    'Limited results to 20 locations max',
    '60-second timeout for external API calls',
    'Optimized bundle size with Vite',
    'MongoDB indexing on frequently queried fields',
    'Lazy loading of components',
    'Image compression before upload'
]
for o in optimizations:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('Error Handling:', 2)
errors = [
    'Graceful fallbacks for external API failures',
    'User-friendly error messages',
    'Warning banners instead of blocking errors',
    'Comprehensive logging with Morgan',
    'Global error handler middleware',
    'Validation at both frontend and backend'
]
for e in errors:
    doc.add_paragraph(e, style='List Bullet')

doc.add_page_break()

# 11. PROJECT STATISTICS
doc.add_heading('11. PROJECT STATISTICS', 1)
stats = [
    'Frontend: 45+ React components',
    'Backend: 4 data models, 8+ API routes',
    'Features: 5 major features (Prescription, Reports, OTC, Disposal, History)',
    'Authentication: 2 methods (OAuth, Email/Password)',
    'Supported Symptoms: 14+ common conditions',
    'Age Groups: 5 categories (infant to senior)',
    'Report Types: 3 main types (blood, radiology, pathology)',
    'Analysis Depths: 3 levels (simple, detailed, educational)',
    'Security Measures: 8+ implemented features'
]
for s in stats:
    doc.add_paragraph(s, style='List Bullet')

# 12. CLOSING POINTS
doc.add_heading('12. CLOSING POINTS', 1)
closing = [
    'HealNet democratizes medical information access',
    'Combines cutting-edge AI with practical healthcare needs',
    'Built with modern, scalable technologies',
    'Prioritizes user privacy and data security',
    'Addresses real-world problems: medication errors, improper disposal',
    'Potential for significant healthcare impact',
    'Ready for production deployment',
    'Extensible architecture for future enhancements'
]
for c in closing:
    doc.add_paragraph(c, style='List Bullet')

# Save document
doc.save('/Users/mayankatri/Desktop/HealNet/HealNet_Presentation_Guide.docx')
print("✅ Document created successfully: HealNet_Presentation_Guide.docx")
